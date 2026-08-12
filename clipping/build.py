"""Gerador de clipping diário no formato Itaú BBA (Word .docx)."""
from __future__ import annotations

import html
import io
import json
import logging
import os
import re
import time
import unicodedata
from dataclasses import dataclass, field
from datetime import date
from html.parser import HTMLParser
from pathlib import Path

log = logging.getLogger(__name__)

_ROOT = Path(__file__).resolve().parent          # pasta do pacote clipping/ (template.docx ao lado)
_TEMPLATE_PATH = _ROOT / "template.docx"
OUT_DIR = _ROOT / "out"


def clipping_basename(d: date) -> str:
    """Nome dos arquivos que saem do clipping: "NEWS - MMDDYYYY" (mês-dia-ano).

    Um lugar só para o Word, o .eml, a prévia e o anexo dentro do e-mail nunca
    divergirem. (Convenção do usuário, 2026-08-11 — antes era clipping_AAAAMMDD.)
    """
    return f"NEWS - {d.strftime('%m%d%Y')}"

# ── Sector detection ──────────────────────────────────────────────────────────

_SM_DOMAINS    = {"core.spglobal.com", "www.mining.com"}
_PP_DOMAINS    = {"dashboard.fastmarkets.com"}


def _nfkd(s: str) -> str:
    return unicodedata.normalize("NFKD", s).lower()


_SM_KWS = {_nfkd(k) for k in [
    "vale", "vale3", "vale s.a.", "vale sa", "vale mining",
    "gerdau", "ggbr", "csn", "usiminas", "ternium", "arcelormittal",
    "metalurgica gerdau", "companhia siderurgica nacional",
    "bhp", "rio tinto", "fortescue", "anglo american",
    "nippon steel", "posco", "tata steel",
    "minerio de ferro", "iron ore", "aco", "steel", "hrc", "slab",
    "billet", "coking coal", "carvao siderurgico",
    "pellet", "pelota", "siderurgia", "steelmaking",
    "mineracao", "copper", "simandou", "usim", "ggbr4",
]}

_PP_KWS = {_nfkd(k) for k in [
    "suzano", "klabin", "cmpc", "veracel", "bracell",
    "eldorado celulose", "empresas copec",
    "celulose", "pulp and paper", "market pulp", "papel e celulose",
    "eucalipto", "bhkp", "bekp", "pulp", "paper", "klbn", "suzb",
]}

def detect_sector(domain: str, matched_keywords: list[str], title: str) -> str:
    """Returns 'SM', 'PP' or 'NR'. (Cement removido — não cobrimos mais.)"""
    if domain in _SM_DOMAINS:
        return "SM"
    if domain in _PP_DOMAINS:
        return "PP"

    kws     = {_nfkd(k) for k in matched_keywords}
    title_n = _nfkd(title)

    sm_score = len(kws & _SM_KWS) + sum(1 for k in _SM_KWS if k in title_n)
    pp_score = len(kws & _PP_KWS) + sum(1 for k in _PP_KWS if k in title_n)

    if pp_score > sm_score:
        return "PP"
    return "SM"


# ── Data model ────────────────────────────────────────────────────────────────

SECTOR_LABEL: dict[str, str] = {
    "NR": "NATURAL RESOURCES",
    "SM": "STEEL & MINING",
    "PP": "PULP & PAPER",
}
SECTOR_ORDER   = ["NR", "SM", "PP"]   # ordem fixa no clipping (usuário 2026-07-27): NR → S&M → P&P
_VALID_SECTORS = frozenset(SECTOR_ORDER)

TAKE_SYMBOL    = {"+": "(+)", "=": "(=)", "-": "(-)"}
TAKE_COLOR_HEX = {"+": "00B050", "-": "FF0000"}   # "=" fica PRETO (sem cor) — igual à referência

# Marcação inline da "Mensagem de abertura": **negrito** e [texto](url).
# grupo 1 = texto em negrito · grupos 2/3 = (texto, url) do link.
_INLINE_RE = re.compile(r'\*\*(.+?)\*\*|\[([^\]]+)\]\((https?://[^)\s]+)\)')


# ══════════════════════════════════════════════════════════════════════════════
# MENSAGEM DE ABERTURA — formatação rica (negrito · itálico · sublinhado · cor · link)
#
# O editor da dashboard grava HTML em config['intro']['html'] (formato canônico, tipo Word).
# Mensagens antigas guardaram só `**negrito**` e `[texto](url)` em config['intro']['text'] —
# o mesmo parser lê as duas. É UMA fonte de verdade: o Word (build.py) e o e-mail (eml.py)
# renderizam a MESMA lista de segmentos, então nunca divergem.
# ══════════════════════════════════════════════════════════════════════════════

@dataclass
class InlineSeg:
    """Trecho contíguo de texto da mensagem de abertura + a formatação dele."""
    text:      str
    bold:      bool = False
    italic:    bool = False
    underline: bool = False
    color:     str  = ""     # 'RRGGBB' (vazio = cor padrão do texto)
    url:       str  = ""     # link externo (vazio = texto comum)


_NAMED_COLORS = {
    "black": "000000", "white": "FFFFFF", "red": "FF0000", "green": "008000",
    "blue": "0000FF", "orange": "FFA500", "yellow": "FFFF00", "purple": "800080",
    "gray": "808080", "grey": "808080", "navy": "000080", "maroon": "800000",
}


def _norm_color(raw: str) -> str:
    """'#f00' | '#FF0000' | 'rgb(255, 0, 0)' | 'red' → 'FF0000'. '' se não reconhecer."""
    v = (raw or "").strip().lower()
    if not v:
        return ""
    m = re.match(r'rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)', v)
    if m:
        return "".join(f"{min(255, int(g)):02X}" for g in m.groups())
    v = v.lstrip("#")
    if re.fullmatch(r'[0-9a-f]{3}', v):
        return "".join(c * 2 for c in v).upper()
    if re.fullmatch(r'[0-9a-f]{6}', v):
        return v.upper()
    return _NAMED_COLORS.get(v, "")


def _safe_url(raw: str) -> str:
    """Só http(s) e mailto entram como link (nada de javascript:/data:)."""
    u = (raw or "").strip()
    return u if re.match(r'^(https?://|mailto:)', u, re.I) else ""


class _IntroHTMLParser(HTMLParser):
    """Percorre o HTML da mensagem de abertura e devolve LINHAS de InlineSeg.

    Whitelist de tags: b/strong · i/em · u/ins · font[color] · span[style=color] ·
    a[href] · br · p/div/li (quebram linha). Tag fora da lista é desembrulhada — o
    conteúdo dela continua sendo lido, só a formatação é ignorada.
    """

    _BLOCKS = {"p", "div", "li", "ul", "ol", "blockquote", "h1", "h2", "h3", "h4"}
    _DROP   = {"script", "style", "head", "title"}
    # tags sem fechamento — não entram na pilha de formatação (senão vazam pro resto do texto)
    _VOID   = {"br", "img", "hr", "input", "meta", "link", "col", "area", "base",
               "source", "embed", "param", "track", "wbr"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.lines: list[list[InlineSeg]] = [[]]
        self._fmt:  list[dict] = []    # pilha de formatação ativa
        self._drop = 0                 # profundidade dentro de tag proibida
        self._brk  = 0                 # quebras de linha pendentes (só valem se vier texto)

    # ── formatação ativa (topo da pilha vence na cor/link) ────────────────────
    def _state(self) -> dict:
        st = {"bold": False, "italic": False, "underline": False, "color": "", "url": ""}
        for f in self._fmt:
            st["bold"]      = st["bold"]      or f.get("bold", False)
            st["italic"]    = st["italic"]    or f.get("italic", False)
            st["underline"] = st["underline"] or f.get("underline", False)
            if f.get("color"):
                st["color"] = f["color"]
            if f.get("url"):
                st["url"] = f["url"]
        return st

    @staticmethod
    def _color_attr(a: dict) -> str:
        """Cor da LETRA da tag: style="color:…" (não 'background-color') ou color="…"."""
        m = re.search(r'(?<!-)\bcolor\s*:\s*([^;]+)', a.get("style", ""), re.I)
        return (_norm_color(m.group(1)) if m else "") or _norm_color(a.get("color", ""))

    def handle_starttag(self, tag, attrs) -> None:
        tag = tag.lower()
        a   = {k.lower(): (v or "") for k, v in attrs}
        if tag in self._DROP:
            self._drop += 1
            return
        if tag == "br":
            self._brk += 1
            return
        if tag in self._BLOCKS and self.lines[-1]:
            self._brk = max(self._brk, 1)   # linha nova só se já houver conteúdo
        if tag in self._VOID:
            return
        fmt: dict = {"_t": tag}
        if tag in ("b", "strong"):
            fmt["bold"] = True
        elif tag in ("i", "em"):
            fmt["italic"] = True
        elif tag in ("u", "ins"):
            fmt["underline"] = True
        elif tag == "a":
            fmt["url"] = _safe_url(a.get("href", ""))
        # a cor pode vir em QUALQUER tag (o editor aproveita o elemento que já envolve a
        # seleção: <i style="color:…">, <a style="color:…">, <div style="color:…">…)
        col = self._color_attr(a)
        if col:
            fmt["color"] = col
        self._fmt.append(fmt)

    def handle_startendtag(self, tag, attrs) -> None:
        if tag.lower() == "br":
            self._brk += 1

    def handle_endtag(self, tag) -> None:
        tag = tag.lower()
        if tag in self._DROP:
            self._drop = max(0, self._drop - 1)
            return
        if tag in self._BLOCKS and self.lines[-1]:
            self._brk = max(self._brk, 1)
        for i in range(len(self._fmt) - 1, -1, -1):     # fecha o mais recente daquela tag
            if self._fmt[i].get("_t") == tag:
                self._fmt.pop(i)
                break

    def handle_data(self, data) -> None:
        if self._drop:
            return
        # HTML colapsa run de espaço/tab/newline em UM espaço. Cuidado: `\s` do Python também
        # casaria o &nbsp; (\xa0) — que é espaço PROPOSITAL do editor e tem que sobreviver.
        text = re.sub(r'[ \t\r\n\f\v]+', ' ', data or "")
        if not text:
            return
        if not text.strip() and not self.lines[-1]:
            return                                        # espaço solto entre tags
        for _ in range(self._brk):                        # materializa as quebras pendentes
            self.lines.append([])
        self._brk = 0
        self.lines[-1].append(InlineSeg(text=text, **self._state()))


def _intro_lines_from_html(src: str) -> list[list[InlineSeg]]:
    p = _IntroHTMLParser()
    try:
        p.feed(src)
        p.close()
    except Exception as e:                                # pragma: no cover
        log.warning("clipping: HTML da mensagem de abertura ilegível (%s) — caindo p/ texto puro", e)
        plain = html.unescape(re.sub(r'<[^>]+>', ' ', src))
        return [[InlineSeg(text=plain.strip())]]
    return p.lines


def _intro_lines_from_markdown(text: str) -> list[list[InlineSeg]]:
    """Formato antigo do editor: `**negrito**` e `[texto](url)`, uma linha por \\n."""
    lines: list[list[InlineSeg]] = []
    for ln in (text or "").splitlines():
        segs: list[InlineSeg] = []
        pos = 0
        for m in _INLINE_RE.finditer(ln):
            if m.start() > pos:
                segs.append(InlineSeg(text=ln[pos:m.start()]))
            if m.group(1) is not None:
                segs.append(InlineSeg(text=m.group(1), bold=True))
            else:
                segs.append(InlineSeg(text=m.group(2), url=_safe_url(m.group(3))))
            pos = m.end()
        if pos < len(ln):
            segs.append(InlineSeg(text=ln[pos:]))
        lines.append(segs)
    return lines


def parse_intro_lines(intro: dict | None) -> list[list[InlineSeg]]:
    """config['intro'] → linhas de segmentos formatados. Prefere o HTML do editor novo;
    cai no markdown antigo (`**bold**`/`[t](u)`) quando não há HTML gravado."""
    intro = intro or {}
    src = (intro.get("html") or "").strip()
    if src:
        return _intro_lines_from_html(src)
    return _intro_lines_from_markdown(intro.get("text") or "")


def line_text(segs: list[InlineSeg]) -> str:
    return "".join(s.text for s in segs)


def intro_has_content(intro: dict | None) -> bool:
    """True se a mensagem de abertura está LIGADA e tem texto de verdade."""
    intro = intro or {}
    if not intro.get("on"):
        return False
    return any(line_text(ln).strip() for ln in parse_intro_lines(intro))


def _is_platts_boilerplate(text: str) -> bool:
    """True se o parágrafo é EXATAMENTE 'Platts is part of S&P Global Energy' (com/sem ponto
    final). Removido a pedido do usuário. Se houver texto DEPOIS (ex.: 'Platts, part of S&P
    Global Energy, assessed…'), NÃO casa → o parágrafo é mantido."""
    t = (text or "").strip().rstrip(".").strip().lower()
    return t == "platts is part of s&p global energy"

# Analistas padrão do bloco de contatos (fallback quando o admin não configura em config['analysts']).
_DEFAULT_ANALYSTS: list[dict] = [
    {"name":  "Daniel Sasson, CFA",
     "role":  "Equity Research – Steel & Mining, Pulp and Paper and Cement",
     "phone": "t. +55 11 3073 3031  m.+55 11 99674 1242",
     "email": "daniel.sasson@itaubba.com"},
    {"name":  "Marcelo Furlan Palhares, CFA",
     "role":  "Equity Research – Steel & Mining, Pulp and Paper and Cement",
     "phone": "t. +55 11 3073 3357  m.+55 11 97464 2801",
     "email": "marcelo.palhares@itaubba.com"},
    {"name":  "João Paulo Luka Helito, CNPI",
     "role":  "Equity Research – Steel & Mining, Pulp and Paper and Cement",
     "phone": "t.+55 11 3073 3005  m.+55 11 93452 7535",
     "email": "joao.helito@itaubba.com"},
]


# Entidade HTML de verdade: &rsquo; &amp; &uacute; &#8217; &#x27; — exige o ";" no fim,
# então um "&" solto ("Pulp & Paper") ou "AT&T" NÃO casa e passa intacto.
_ENTITY_RE = re.compile(r"&(?:[a-zA-Z][a-zA-Z0-9]{1,31}|#\d{1,7}|#[xX][0-9a-fA-F]{1,6});")


def clean_headline(t: str) -> str:
    """Manchete → texto de verdade. "Arauco&rsquo;s Sucuri&uacute;" → "Arauco’s Sucuriú".

    O Fastmarkets entrega o título com CÓDIGO DE HTML dentro e o banco guarda assim (122 de
    124 casos medidos em 11.540 manchetes). Na dashboard o navegador decodifica sozinho e
    ninguém vê; no Word/e-mail vira texto literal — era o título "todo cagado" no clipping.

    Rede de segurança do lado do clipping: o `hunter/fastmarkets_scraper.py` já decodifica na
    coleta, mas as manchetes que JÁ estão no banco vieram cruas (o push é ignore-duplicates,
    não reescreve título). Só age quando existe entidade de fato — medido: dos 11.416 títulos
    sem entidade, ZERO mudariam. Um passe basta (0 casos de codificação dupla no banco).
    """
    if not t or "&" not in t:
        return t
    return html.unescape(t) if _ENTITY_RE.search(t) else t


@dataclass
class ClippingItem:
    url:              str
    title:            str
    source_name:      str
    body:             str
    matched_keywords: list[str]
    domain:           str
    take:             str            # '+', '=' or '-'
    sector:           str = field(default="SM")
    # Campos opcionais para artigos bilíngues (Valor/Estadão/El Financiero)
    translated_title: str = field(default="")
    translated_body:  str = field(default="")

    def __post_init__(self):
        # vale p/ QUEM construir o item (geração, aquecedor, testes) — não dá p/ esquecer
        self.title            = clean_headline(self.title)
        self.translated_title = clean_headline(self.translated_title)


# ── HTML → structured blocks (texto + imagens) ────────────────────────────────

def _html_to_blocks(html_body: str) -> list[dict]:
    """Converte HTML do corpo do artigo em blocos estruturados.

    Retorna lista de dicts (espelhando a hierarquia visual do leitor da dashboard):
      {"type": "text",       "text": str, "bold": bool}
      {"type": "h3",         "text": str}          ← subtítulo de seção
      {"type": "h4",         "text": str}          ← subtítulo menor
      {"type": "blockquote", "text": str}          ← citação recuada
      {"type": "list_item",  "text": str}
      {"type": "image",      "src": str}
    """
    if not html_body:
        return []
    try:
        from bs4 import BeautifulSoup, NavigableString, Tag
    except ImportError:
        # Fallback simples sem BS4
        text = re.sub(r"<[^>]+>", " ", html_body)
        import html as _h
        text = _h.unescape(text)
        return [{"type": "text", "text": p.strip(), "bold": False}
                for p in re.split(r"\n+", text) if len(p.strip()) > 5]

    soup = BeautifulSoup(html_body, "lxml")
    blocks: list[dict] = []

    def _clean(text: str) -> str:
        # Remove caracteres de substituição Unicode (U+FFFD) de encoding errado
        return text.replace("�", "").strip()

    def _add_text(text: str, bold: bool = False) -> None:
        text = _clean(text)
        if _is_platts_boilerplate(text):   # remove o boilerplate solto "Platts is part of S&P Global Energy"
            return
        if text and len(text) > 3:
            blocks.append({"type": "text", "text": text, "bold": bold})

    def _process(el) -> None:
        if isinstance(el, NavigableString):
            t = _clean(str(el))
            if t:
                _add_text(t)
            return
        if not isinstance(el, Tag):
            return

        name = el.name

        if name == "img":
            src = (el.get("src") or "").strip()
            if src and not src.startswith("data:image/gif"):  # skip tracking pixels
                blocks.append({"type": "image", "src": src})

        elif name in ("h1", "h2"):
            # h1/h2 tratados como subtítulo h3 (não sobrepõem o título do artigo)
            text = _clean(el.get_text(" ", strip=True))
            if text:
                blocks.append({"type": "h3", "text": text})

        elif name == "h3":
            text = _clean(el.get_text(" ", strip=True))
            if text:
                blocks.append({"type": "h3", "text": text})

        elif name in ("h4", "h5", "h6"):
            text = _clean(el.get_text(" ", strip=True))
            if text:
                blocks.append({"type": "h4", "text": text})

        elif name == "blockquote":
            text = _clean(el.get_text(" ", strip=True))
            if text:
                blocks.append({"type": "blockquote", "text": text})

        elif name in ("p", "div", "section", "article"):
            # Extrair imagens internas antes de pegar o texto
            for img in el.find_all("img"):
                src = (img.get("src") or "").strip()
                if src and not src.startswith("data:image/gif"):
                    blocks.append({"type": "image", "src": src})
                img.decompose()
            text = _clean(el.get_text(" ", strip=True))
            _add_text(text)

        elif name in ("ul", "ol"):
            for li in el.find_all("li", recursive=False):
                text = _clean(li.get_text(" ", strip=True))
                if text:
                    blocks.append({"type": "list_item", "text": text})

        elif name == "li":
            text = _clean(el.get_text(" ", strip=True))
            if text:
                blocks.append({"type": "list_item", "text": text})

        elif name in ("strong", "b"):
            _add_text(el.get_text(" ", strip=True), bold=True)

        elif name == "br":
            pass

        else:
            for child in el.children:
                _process(child)

    body_el = soup.find("body") or soup
    for child in body_el.children:
        _process(child)

    # Se nenhum bloco foi encontrado, usa fallback plano
    if not blocks:
        for el in soup.find_all(["p", "h2", "h3", "img", "li"]):
            _process(el)

    return blocks


# ── Domínios que NÃO devem ser auto-fetchados (têm scrapers dedicados) ────────

_SKIP_AUTO_FETCH: frozenset[str] = frozenset([
    "core.spglobal.com",
    "dashboard.fastmarkets.com",
    "valor.globo.com",
    "www.estadao.com.br",
    "www.worldcement.com",   # corpo já capturado pelo worldcement_scraper
])

# ── Domínios que geram clipping bilíngue (original + Free Translation) ────────

_BILINGUAL_DOMAINS: frozenset[str] = frozenset([
    "valor.globo.com",          # Português → Inglês
    "www.estadao.com.br",       # Português → Inglês
    "portalcelulose.com.br",    # Português → Inglês (2026-08-03, pedido do usuário)
    "www.elfinanciero.com.mx",  # Espanhol → Inglês
])

_DOMAIN_LANG: dict[str, str] = {
    "valor.globo.com":          "Portuguese",
    "www.estadao.com.br":       "Portuguese",
    "portalcelulose.com.br":    "Portuguese",
    "www.elfinanciero.com.mx":  "Spanish",
}


def _fetch_body_regular(url: str) -> str:
    """Busca corpo completo de um artigo em sites regulares (sem paywall/SPA).

    Usado pelo gerador de clipping quando o artigo foi indexado sem corpo
    (ex.: mining.com, el financiero) e o usuário não o abriu no leitor antes.

    Extrai o container <article> / <main> via requests + BeautifulSoup e
    sanitiza com article_to_safe_html para preservar estrutura (h2, ul, img).
    """
    try:
        import requests as _req
        from bs4 import BeautifulSoup as _BS
        from .html_utils import article_to_safe_html, extract_article_container

        html = ""
        try:   # curl_cffi (impersonate) — passa em sites que bloqueiam requests puro (ex.: Portal Celulose)
            from curl_cffi import requests as _creq
            r = _creq.get(url, impersonate="chrome124", timeout=20, allow_redirects=True)
            if r.status_code == 200:
                html = r.text
        except Exception:
            html = ""
        if not html:   # fallback requests puro
            resp = _req.get(
                url,
                headers={"User-Agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/124.0.0.0 Safari/537.36"
                )},
                timeout=15,
                allow_redirects=True,
            )
            if resp.status_code != 200:
                log.debug("clipping: auto-fetch HTTP %d em %s", resp.status_code, url)
                return ""
            html = resp.text

        soup = _BS(html, "lxml")

        # Extrai o container ANTES de limpar (o decompose global removia o <article> em alguns
        # temas WordPress, ex.: Portal Celulose) — depois limpa o ruído só DENTRO do container.
        container = extract_article_container(soup, url)
        if not container:
            for _sel in (".entry-content", ".post-content", ".td-post-content",
                         "article", "[role=main]", "main"):
                _c = soup.select_one(_sel)
                if _c and len(_c.get_text(strip=True)) > 200:
                    container = _c
                    break
        if not container:
            return ""
        for tag in container.find_all([
            "nav", "header", "footer", "aside",
            "script", "style", "iframe", "noscript", "form",
        ]):
            tag.decompose()
        for tag in container.find_all(class_=re.compile(
            r"(^|\b)(ad|ads|banner|popup|cookie|subscribe|paywall|"
            r"menu|sidebar|share|social|comment|promo|newsletter)(\b|$)",
            re.I,
        )):
            tag.decompose()

        body = article_to_safe_html(str(container))
        if body and len(body) > 100:
            return body

    except Exception as e:
        log.debug("clipping: auto-fetch falhou em %s: %s", url, e)

    return ""


# ── Download de imagem ────────────────────────────────────────────────────────

def _fetch_image(src: str, article_url: str = "") -> tuple[bytes, str] | None:
    """Baixa imagem de src. Retorna (dados_bytes, extensão) ou None se falhar."""
    import base64

    # Data URI (base64 inline)
    if src.startswith("data:"):
        try:
            header, data = src.split(",", 1)
            mime = header.split(";")[0].split(":")[1]   # ex: "image/png"
            ext  = mime.split("/")[1]                   # ex: "png"
            if ext == "jpeg":
                ext = "jpg"
            if ext in ("png", "jpg", "gif"):
                return base64.b64decode(data), ext
        except Exception:
            pass
        return None

    # Resolve URL relativa
    if src.startswith("//"):
        src = "https:" + src
    elif src.startswith("/") and article_url:
        from urllib.parse import urlparse
        p = urlparse(article_url)
        src = f"{p.scheme}://{p.netloc}{src}"

    if not src.startswith("http"):
        return None

    try:
        import requests
        r = requests.get(
            src,
            headers={"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"},
            timeout=10,
            allow_redirects=True,
        )
        if r.status_code != 200 or len(r.content) < 500:
            return None
        ct   = r.headers.get("Content-Type", "image/jpeg").split(";")[0].strip()
        ext  = {"image/jpeg": "jpg", "image/png": "png", "image/gif": "gif",
                "image/webp": "webp"}.get(ct, "jpg")
        # Tenta converter WebP para JPEG via Pillow (opcional)
        if ext == "webp":
            try:
                from PIL import Image as _PILImage
                img_pil = _PILImage.open(io.BytesIO(r.content)).convert("RGB")
                buf = io.BytesIO()
                img_pil.save(buf, format="JPEG", quality=85)
                return buf.getvalue(), "jpg"
            except Exception:
                return None   # WebP sem Pillow → pula
        return r.content, ext
    except Exception:
        return None


# ── Tradução bilíngue ─────────────────────────────────────────────────────────

_LANG_CODE: dict[str, str] = {
    "Portuguese": "pt",
    "Spanish":    "es",
}


# Pacing GLOBAL da tradução: garante >= _PACE_S entre quaisquer duas chamadas ao Google
# Translate grátis (na geração inteira, todos os artigos) — PREVINE o rate-limit em vez de
# só reagir a ele. Era a causa de "1 de 2 não traduziu": rajada de chamadas → bloqueio.
_LAST_TRANSLATE = [0.0]
_PACE_S = 1.2


def _translate_chain() -> list[dict]:
    """Provedores de TRADUÇÃO do clipping, EM ORDEM.

    2026-08-06 — DESACOPLADO do Gemini. Antes a tradução usava `CLIPPING_TRANSLATE_KEY` **ou,
    se ausente, `GEMINI_API_KEY`** — a MESMA cota dos takes. Desde `1a0ae81` o Gemini é o
    TITULAR dos takes, e o clipping roda no FIM do dia → era ele quem ficaria sem cota, na
    hora da entrega. Agora:

      1) gemini DEDICADO (`CLIPPING_TRANSLATE_KEY`) — chave de um PROJETO Google SEPARADO só
                      p/ o clipping (2026-08-07: projeto "clipping" da separação por carga).
                      Cota isolada de 500/dia p/ um consumo de ~7/dia → folga absurda, e é o
                      melhor modelo. ⚠️ Chave nova no MESMO projeto NÃO adianta: a cota grátis
                      do Gemini é POR PROJETO.
      2) zai (GLM)  — fallback. Traduzir é a tarefa FÁCIL (o GLM só se mostrou fraco em
                      JULGAMENTO de take, não em língua) e ele tem ~1.000/dia ociosos.
      3) gemini COMPARTILHADO (`GEMINI_API_KEY`) — ÚLTIMO recurso, e de propósito: essa é a
                      cota dos TAKES. São ~7 chamadas por clipping (medido), dano desprezível,
                      mas só se os dois de cima falharem.

    SEM o secret dedicado a cadeia vira `zai -> gemini compartilhado` = exatamente o
    comportamento de hoje; nada quebra enquanto o projeto novo não existir.
    Depois desta cadeia, o chamador ainda cai no Google Translate grátis.
    Ordem trocável por CLIPPING_TRANSLATE_CHAIN (ex.: "zai,gemini_dedicated").
    """
    try:
        from hunter.llm_take import PROVIDERS as _P
        gem_model = (_P.get("gemini") or {}).get("model") or "gemini-2.5-flash-lite"
        zai_model = (_P.get("zai") or {}).get("model") or "glm-4.5-flash"
    except Exception:
        gem_model, zai_model = "gemini-2.5-flash-lite", "glm-4.5-flash"
    gem_model = os.environ.get("CLIPPING_TRANSLATE_MODEL", gem_model)

    defs = {
        "zai": {"name": "zai", "kind": "openai", "key": os.environ.get("ZAI_API_KEY", ""),
                "model": os.environ.get("CLIPPING_TRANSLATE_ZAI_MODEL", zai_model),
                "url": "https://api.z.ai/api/paas/v4/chat/completions"},
        "gemini_dedicated": {"name": "gemini(dedicada)", "kind": "gemini",
                             "key": os.environ.get("CLIPPING_TRANSLATE_KEY", ""), "model": gem_model},
        "gemini_shared": {"name": "gemini(compartilhada)", "kind": "gemini",
                          "key": os.environ.get("GEMINI_API_KEY", ""), "model": gem_model},
    }
    order = [p.strip() for p in os.environ.get(
        "CLIPPING_TRANSLATE_CHAIN", "gemini_dedicated,zai,gemini_shared").split(",") if p.strip()]
    return [defs[p] for p in order if p in defs and defs[p]["key"]]


def _translate_call(cfg: dict, title: str, para_texts: list[str], source_lang: str):
    """Uma chamada de tradução. Devolve (title, paragraphs) ou None."""
    import requests
    system = (f"You are a professional news translator. Translate from {source_lang} to English, "
              "preserving meaning, journalistic tone, proper names and numbers. Do not summarize "
              "or add commentary.")
    user = ("Translate the article to English. Keep the SAME number of paragraphs, in order. "
            'Reply ONLY with JSON: {"title": "...", "paragraphs": ["...", "..."]}\n\n'
            + json.dumps({"title": title, "paragraphs": para_texts}, ensure_ascii=False))

    if cfg["kind"] == "gemini":
        url = f"https://generativelanguage.googleapis.com/v1beta/models/{cfg['model']}:generateContent"
        payload = {
            "systemInstruction": {"parts": [{"text": system}]},
            "contents": [{"role": "user", "parts": [{"text": user}]}],
            "generationConfig": {
                "temperature": 0, "maxOutputTokens": 8192,
                "responseMimeType": "application/json",
                "responseSchema": {"type": "OBJECT", "properties": {
                    "title": {"type": "STRING"},
                    "paragraphs": {"type": "ARRAY", "items": {"type": "STRING"}}},
                    "required": ["title", "paragraphs"]}},
        }
        r = requests.post(url, json=payload,
                          headers={"x-goog-api-key": cfg["key"], "Content-Type": "application/json"},
                          timeout=90)
        if not r.ok:
            log.warning("clipping: tradução %s HTTP %s", cfg["name"], r.status_code)
            return None
        raw = r.json()["candidates"][0]["content"]["parts"][0]["text"]
    else:                                            # OpenAI-compatível (Z.AI)
        payload = {"model": cfg["model"], "temperature": 0,
                   "response_format": {"type": "json_object"},   # GLM sem isto às vezes não devolve JSON limpo
                   "messages": [{"role": "system", "content": system},
                                {"role": "user", "content": user}]}
        r = requests.post(cfg["url"], json=payload,
                          headers={"Authorization": f"Bearer {cfg['key']}",
                                   "Content-Type": "application/json"}, timeout=90)
        if not r.ok:
            log.warning("clipping: tradução %s HTTP %s", cfg["name"], r.status_code)
            return None
        raw = r.json()["choices"][0]["message"]["content"]

    out = None
    try:
        out = json.loads(raw)
    except Exception:                                # modelos "thinking" (GLM) podem embrulhar o JSON
        try:
            from hunter.llm_take import _json_candidates
            for cand in _json_candidates(raw):
                if isinstance(cand, dict) and cand.get("title"):
                    out = cand
                    break
        except Exception:
            out = None
    if not isinstance(out, dict):
        log.warning("clipping: tradução %s devolveu JSON inválido", cfg["name"])
        return None
    t_title = (out.get("title") or "").strip()
    t_paras = [str(p).strip() for p in (out.get("paragraphs") or []) if str(p).strip()]
    return (t_title, t_paras) if t_title else None


def _translate_via_llm(title: str, para_texts: list[str], source_lang: str) -> tuple[str, list[str]] | None:
    """Traduz título + parágrafos p/ inglês via IA: 1 chamada, SEM o rate-limit do Google
    Translate grátis. Percorre a cadeia de _translate_chain() (ver lá o porquê da ordem).
    Retorna None se todos falharem (aí o chamador cai no Google Translate)."""
    chain = _translate_chain()
    if not chain:
        return None
    for cfg in chain:
        try:
            res = _translate_call(cfg, title, para_texts, source_lang)
        except Exception as e:
            log.warning("clipping: tradução %s falhou (%s)", cfg["name"], e)
            continue
        if res:
            log.info("clipping: tradução via %s (%s)", cfg["name"], cfg["model"])
            return res
    log.warning("clipping: toda a cadeia de tradução falhou — cai p/ Google")
    return None


def _translate_to_english(title: str, body_html: str, source_lang: str) -> tuple[str, str]:
    """Traduz título + corpo p/ inglês. PRIMÁRIO: IA (Gemini, 1 chamada, confiável, SEM
    rate-limit). FALLBACK: Google Translate grátis (lote + pacing + retry). ('','') se ambos
    falharem → o item fica monolíngue (sem lixo)."""
    try:
        from bs4 import BeautifulSoup as _BS
    except ImportError as e:
        log.warning("clipping: bs4 ausente (%s) — sem tradução", e)
        return "", ""

    # Textos-folha do corpo (parágrafos/headings/itens), NA ORDEM — usados pela IA e pelo Google.
    para_texts: list[str] = []
    if body_html:
        _root = _BS(body_html, "lxml")
        _root = _root.find("body") or _root
        for tag in _root.find_all(["p", "h2", "h3", "h4", "li", "blockquote"]):
            if tag.find(["p", "h2", "h3", "h4", "li", "blockquote"]):     # pula container aninhado
                continue
            t = tag.get_text(separator=" ", strip=True)
            if t:
                para_texts.append(t)

    def _paras_to_html(paras: list[str]) -> str:
        return "\n".join(f"<p>{html.escape(p.strip())}</p>" for p in paras if p.strip())

    # ── 1) IA (confiável, sem rate-limit) ──────────────────────────────────────
    _llm = _translate_via_llm(title, para_texts, source_lang)
    if _llm:
        t_title, t_paras = _llm
        log.info("clipping: tradução via IA '%s' → '%s'", title[:45], t_title[:45])
        return t_title, _paras_to_html(t_paras)

    # ── 2) FALLBACK: Google Translate grátis (lote + pacing + retry) ───────────
    try:
        from deep_translator import GoogleTranslator as _GT
    except ImportError as e:
        log.warning("clipping: sem IA (sem chave) e sem deep_translator (%s) — item monolíngue", e)
        return "", ""
    lang_code = _LANG_CODE.get(source_lang, "pt")

    def _translate_text(text: str) -> str:
        """Bloco (chunks de 4500) com PACING (>=_PACE_S entre chamadas) + retry com backoff
        p/ o rate-limit do Google grátis. Só devolve o original se todas as tentativas falharem."""
        if not text or not text.strip():
            return text
        out: list[str] = []
        for i in range(0, len(text), 4500):
            chunk = text[i:i + 4500]
            translated = None
            for attempt in range(5):
                _wait = _PACE_S - (time.time() - _LAST_TRANSLATE[0])
                if _wait > 0:
                    time.sleep(_wait)
                try:
                    r = _GT(source=lang_code, target="en").translate(chunk)
                    _LAST_TRANSLATE[0] = time.time()
                    if r and r.strip():
                        translated = r
                        break
                except Exception as e:
                    _LAST_TRANSLATE[0] = time.time()
                    log.debug("clipping: Google tradução falhou (%d/5): %s", attempt + 1, e)
                if attempt < 4:
                    time.sleep(2 * (2 ** attempt))   # backoff: 2s, 4s, 8s, 16s
            out.append(translated or chunk)
        return " ".join(out)

    def _translate_batch(texts: list[str]) -> list[str]:
        """Agrupa em lotes < 4000 chars e traduz cada um de uma vez (parágrafos por \\n),
        separando de volta. Fallback item-a-item se a contagem de linhas não bater."""
        result: list[str] = []
        i, n = 0, len(texts)
        while i < n:
            batch, blen = [], 0
            while i < n and (not batch or blen + len(texts[i]) < 4000):
                batch.append(texts[i]); blen += len(texts[i]) + 1; i += 1
            parts = _translate_text("\n".join(t.replace("\n", " ") for t in batch)).split("\n")
            if len(parts) == len(batch):
                result.extend(p.strip() for p in parts)
            else:
                result.extend(_translate_text(t) for t in batch)
        return result

    try:
        t_title = _translate_text(title)
        if not t_title or t_title == title:
            log.warning("clipping: Google tradução do título sem efeito — '%s'", title[:50])
            return "", ""
        t_body = _paras_to_html(_translate_batch(para_texts)) if para_texts else ""
        return t_title, t_body
    except Exception as e:
        log.warning("clipping: Google tradução falhou p/ '%s': %s", title[:50], e)
        return "", ""


# ── Word document builder ─────────────────────────────────────────────────────

def _update_banner_date(body_el, d: date) -> None:
    """Atualiza a data no text box flutuante do banner (parágrafo 0).

    O template tem a data dividida em runs individuais:
      '0' '5' '/' 'XX'(yellow) '/202' '6'
    Este helper substitui os dígitos pelo dia/mês/ano real e remove o highlight.
    """
    WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    tag_r   = f"{{{WNS}}}r"
    tag_t   = f"{{{WNS}}}t"
    tag_rPr = f"{{{WNS}}}rPr"
    tag_hl  = f"{{{WNS}}}highlight"

    first_p = list(body_el)[0]   # Parágrafo 0 = banner com shape flutuante

    # Coleta todos os <w:r> com texto no banner
    runs = []
    for r_el in first_p.iter(tag_r):
        t_el   = r_el.find(tag_t)
        rPr_el = r_el.find(tag_rPr)
        hl_el  = rPr_el.find(tag_hl) if rPr_el is not None else None
        runs.append({
            "r": r_el, "t": t_el,
            "text": (t_el.text or "") if t_el is not None else "",
            "hl":   (hl_el.get(f"{{{WNS}}}val") if hl_el is not None else ""),
            "hl_el": hl_el, "rPr": rPr_el,
        })

    month_str = f"{d.month:02d}"   # "05"
    day_str   = f"{d.day:02d}"     # "19"
    year_str  = str(d.year)        # "2026"

    for i, run in enumerate(runs):
        if run["hl"] == "yellow" and run["text"] == "XX":
            # Substitui dia e remove highlight amarelo
            if run["t"] is not None:
                run["t"].text = day_str
            if run["hl_el"] is not None and run["rPr"] is not None:
                run["rPr"].remove(run["hl_el"])

            # Mês: runs em i-3 ("0") e i-2 ("5")
            try:
                if runs[i-3]["t"] is not None:
                    runs[i-3]["t"].text = month_str[0]
                if runs[i-2]["t"] is not None:
                    runs[i-2]["t"].text = month_str[1]
            except (IndexError, KeyError):
                pass

            # Ano: runs em i+1 ("/202") e i+2 ("6")
            try:
                if runs[i+1]["t"] is not None:
                    runs[i+1]["t"].text = "/" + year_str[:3]
                if runs[i+2]["t"] is not None:
                    runs[i+2]["t"].text = year_str[3]
            except (IndexError, KeyError):
                pass


def _build_word(items: list[ClippingItem], d: date, config: dict | None = None) -> bytes:
    config = config or {}
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    FONT = "Arial"
    # Largura do conteúdo A5 com as margens do template (6237 DXA = 4.33 in)
    CONTENT_WIDTH = Inches(4.33)

    # Abre template → herda logo no header, estilos, numeração e tamanho de página
    doc = Document(str(_TEMPLATE_PATH)) if _TEMPLATE_PATH.exists() else Document()

    # Neutraliza o estilo "Hyperlink" para que os links internos (âncora) não
    # apareçam em azul/sublinhado — Word aplica esse estilo automaticamente a
    # qualquer elemento w:hyperlink, inclusive os de âncora interna.
    _WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    _styles_el = doc.part.styles._element
    for _style_el in _styles_el.findall(f"{{{_WNS}}}style"):
        _name_el = _style_el.find(f"{{{_WNS}}}name")
        if _name_el is not None and _name_el.get(f"{{{_WNS}}}val") == "Hyperlink":
            _rPr_el = _style_el.find(f"{{{_WNS}}}rPr")
            if _rPr_el is not None:
                for _color_el in list(_rPr_el.findall(f"{{{_WNS}}}color")):
                    _rPr_el.remove(_color_el)
                for _u_el in list(_rPr_el.findall(f"{{{_WNS}}}u")):
                    _rPr_el.remove(_u_el)
            break

    # Mantém os 3 primeiros parágrafos do template:
    #   [0] Parágrafo com text box flutuante (banner "Itaú BBA | Equity Research + data")
    #   [1] Parágrafo vazio (espaçador)
    #   [2] Parágrafo com logo inline
    # Remove todo o restante (exceto sectPr).
    body = doc.element.body
    children = list(body)
    for child in children[3:]:
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    # Atualiza a data no banner com o dia/mês/ano real
    _update_banner_date(body, d)

    # ── Helpers ───────────────────────────────────────────────────────────────

    def _zero_spacing(para):
        pPr = para._p.get_or_add_pPr()
        sp  = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0")
        sp.set(qn("w:beforeAutospacing"), "0")
        sp.set(qn("w:after"),  "0")
        sp.set(qn("w:afterAutospacing"), "0")
        pPr.append(sp)

    def _justify(para):
        pPr = para._p.get_or_add_pPr()
        jc  = OxmlElement("w:jc")
        jc.set(qn("w:val"), "both")
        pPr.append(jc)

    def _para_shading(para, fill_hex: str):
        """Fundo de parágrafo inteiro (ex: azul-marinho #000080)."""
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  fill_hex)
        pPr.append(shd)

    def _add_numPr(para, num_id: int = 1, ilvl: int = 0):
        pPr   = para._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl_el = OxmlElement("w:ilvl")
        ilvl_el.set(qn("w:val"), str(ilvl))
        numPr.append(ilvl_el)
        numId_el = OxmlElement("w:numId")
        numId_el.set(qn("w:val"), str(num_id))
        numPr.append(numId_el)
        pPr.append(numPr)

    def _run(para, text, *, bold=False, italic=False, underline=False, size_pt=None,
             color_hex=None, hl=None):
        from docx.shared import RGBColor
        run = para.add_run(text)
        run.font.name   = FONT
        run.font.bold   = bold
        run.font.italic = italic
        if underline:
            run.font.underline = True
        if size_pt is not None:
            run.font.size = Pt(size_pt)
        if color_hex:
            r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        if hl:
            rPr   = run._r.get_or_add_rPr()
            hl_el = OxmlElement("w:highlight")
            hl_el.set(qn("w:val"), hl)
            rPr.append(hl_el)
        return run

    def _heading_index(text: str) -> None:
        """Título de seção do índice: 'Sector Headlines' / 'Recent Publications'.
        Fundo preto, texto branco — mesma linguagem visual dos cabeçalhos de setor."""
        p = doc.add_paragraph()
        _zero_spacing(p)
        _justify(p)
        _run(p, text, bold=True, size_pt=16, color_hex="FFFFFF", hl="black")  # highlight preto (largura do texto)

    def _heading_sector(text: str) -> None:
        """Barra preta: STEEL & MINING / PULP & PAPER / etc."""
        p = doc.add_paragraph()
        _zero_spacing(p)
        _justify(p)
        _run(p, text, bold=True, size_pt=16, color_hex="FFFFFF", hl="black")  # highlight preto (largura do texto)

    def _blank() -> None:
        p = doc.add_paragraph()
        _zero_spacing(p)

    def _blanks(n: int = 1) -> None:
        """Insere n linhas em branco (parágrafos vazios). Unidade de espaçamento entre
        seções — a marca de parágrafo é fixada em Arial 11 p/ altura de linha previsível."""
        for _ in range(max(0, n)):
            p = doc.add_paragraph()
            _zero_spacing(p)
            pPr = p._p.get_or_add_pPr()
            rPr = OxmlElement("w:rPr")
            rf  = OxmlElement("w:rFonts"); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)
            rPr.append(rf)
            sz  = OxmlElement("w:sz"); sz.set(qn("w:val"), "22"); rPr.append(sz)   # 11pt
            pPr.append(rPr)

    def _intro_para_fmt(para) -> None:
        """Formato do parágrafo da Mensagem de abertura (referência): justificado,
        entrelinha 'ao menos 16pt' (w:line=320, atLeast), sem espaço antes/depois."""
        pPr = para._p.get_or_add_pPr()
        sp  = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "0"); sp.set(qn("w:beforeAutospacing"), "0")
        sp.set(qn("w:after"),  "0"); sp.set(qn("w:afterAutospacing"),  "0")
        sp.set(qn("w:line"), "320"); sp.set(qn("w:lineRule"), "atLeast")   # 16pt = 320 twips
        pPr.append(sp)
        _justify(para)

    # Bookmarks para hyperlinks internos.
    # Inicia em 100 para evitar colisão com IDs já presentes no template.docx.
    bm_id_counter = [100]

    def _add_bookmark(para, bm_name: str) -> None:
        """Insere bookmarkStart + bookmarkEnd envolvendo os runs já existentes.

        IMPORTANTE: deve ser chamado APÓS o(s) _run() do parágrafo, para que o
        bookmark abraçe o texto (bm_start antes do primeiro run, bm_end no final).
        Se chamado com parágrafo vazio o bookmark ficaria vazio e alguns clientes
        Word não conseguem navegar até ele de forma confiável.
        """
        bm_id    = bm_id_counter[0]
        bm_id_counter[0] += 1
        p        = para._p
        bm_start = OxmlElement("w:bookmarkStart")
        bm_start.set(qn("w:id"),   str(bm_id))
        bm_start.set(qn("w:name"), bm_name)
        bm_end   = OxmlElement("w:bookmarkEnd")
        bm_end.set(qn("w:id"),     str(bm_id))
        # bm_start: logo após pPr (antes de qualquer run já adicionado)
        pPr = p.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(bm_start)
        else:
            p.insert(0, bm_start)
        # bm_end: no final do parágrafo (após todos os runs)
        p.append(bm_end)

    def _external_hyperlink_run(para, text: str, url: str, *, size_pt: float = 9.5,
                                color: str = "000000", underline: bool = False,
                                bold: bool = False, italic: bool = False) -> None:
        """Run com hyperlink externo (URL) — abre no browser. color/underline configuráveis:
        publicações = preto liso (padrão); link na mensagem de abertura = azul sublinhado."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        rId = para.part.relate_to(url, RT.HYPERLINK, is_external=True)
        hl_el = OxmlElement("w:hyperlink")
        hl_el.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", rId)
        r    = OxmlElement("w:r")
        rPr  = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT)
        rFonts.set(qn("w:hAnsi"), FONT)
        rPr.append(rFonts)
        if bold:
            rPr.append(OxmlElement("w:b"))
        if italic:
            rPr.append(OxmlElement("w:i"))
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(int(size_pt * 2)))   # Word: 2 × pt
        rPr.append(sz)
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), color)
        rPr.append(color_el)
        if underline:
            u_el = OxmlElement("w:u"); u_el.set(qn("w:val"), "single"); rPr.append(u_el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        hl_el.append(r)
        para._p.append(hl_el)

    def _mailto_run(para, email: str, *, size_pt: float = 10) -> None:
        """E-mail como hyperlink AZUL SUBLINHADO (mailto) — igual à referência."""
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        rId = para.part.relate_to(f"mailto:{email}", RT.HYPERLINK, is_external=True)
        hl_el = OxmlElement("w:hyperlink")
        hl_el.set("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id", rId)
        r   = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts"); rFonts.set(qn("w:ascii"), FONT); rFonts.set(qn("w:hAnsi"), FONT)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz"); sz.set(qn("w:val"), str(int(size_pt * 2))); rPr.append(sz)
        color_el = OxmlElement("w:color"); color_el.set(qn("w:val"), "0000FF"); rPr.append(color_el)
        u_el = OxmlElement("w:u"); u_el.set(qn("w:val"), "single"); rPr.append(u_el)
        r.append(rPr)
        t = OxmlElement("w:t"); t.text = email
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        hl_el.append(r)
        para._p.append(hl_el)

    def _hyperlink_run(para, text: str, bm_name: str) -> None:
        """Run com hyperlink interno (âncora) visível no sumário.

        Estilo: sublinhado preto — indica ao leitor que o título é clicável
        (Ctrl+Click no Word; clique simples no PDF exportado).
        Só o texto do título recebe o link — o rótulo do setor e a fonte não.
        """
        hl_el = OxmlElement("w:hyperlink")
        hl_el.set(qn("w:anchor"), bm_name)
        r    = OxmlElement("w:r")
        rPr  = OxmlElement("w:rPr")
        # Fonte
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), FONT)
        rFonts.set(qn("w:hAnsi"), FONT)
        rPr.append(rFonts)
        # Tamanho 11pt (igual aos outros runs do bullet)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "22")
        rPr.append(sz)
        # Cor: AZUL de link (2026-08-03, pedido do usuário — "corzinha azul p/ visualizar").
        # Sobrescreve o azul automático do estilo Hyperlink (neutralizado) com o mesmo 0000FF.
        color_el = OxmlElement("w:color")
        color_el.set(qn("w:val"), "0000FF")
        rPr.append(color_el)
        r.append(rPr)
        t = OxmlElement("w:t")
        t.text = text
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r.append(t)
        hl_el.append(r)
        para._p.append(hl_el)

    def _add_image_para(src: str, article_url: str) -> None:
        """Baixa e insere imagem centralizada. Espelha .reader-img (max-width, margin auto).
        Pula silenciosamente se falhar."""
        img_result = _fetch_image(src, article_url)
        if not img_result:
            return
        data, _ext = img_result
        try:
            img_stream = io.BytesIO(data)
            p_img = doc.add_paragraph()
            # Espaçamento: espelha margin: 1.2rem auto do CSS
            p_img.paragraph_format.space_before = Pt(6)
            p_img.paragraph_format.space_after  = Pt(6)
            p_img.paragraph_format.alignment    = 1   # WD_ALIGN_PARAGRAPH.CENTER
            run = p_img.add_run()
            run.add_picture(img_stream, width=CONTENT_WIDTH)
        except Exception as e:
            log.debug("clipping: imagem não inserida (%s): %s", src[:80], e)

    # ── Ordem dos setores ─────────────────────────────────────────────────────
    # Ordem dos setores no clipping = ordem CANÔNICA NR → SM → PP (SECTOR_ORDER), NUNCA a
    # ordem do payload → S&M SEMPRE antes de P&P (exigência do usuário). Só entra setor presente.
    seen_sectors: list[str] = [s for s in SECTOR_ORDER if any(it.sector == s for it in items)]

    bm_names    = {item.url: f"art{i}"   for i, item in enumerate(items)}
    bm_tr_names = {item.url: f"art{i}tr" for i, item in enumerate(items)}   # âncora da Free Translation

    # ══════════════════════════════════════════════════════════════════════════
    # INTRO / MENSAGEM (configurável — vai no topo; o mesmo texto entra no e-mail)
    # ══════════════════════════════════════════════════════════════════════════
    _intro = (config.get("intro") or {})
    _has_intro = intro_has_content(_intro)
    if _has_intro:
        _blanks(1)   # logo do Itaú → Mensagem de abertura: 1 linha
        # negrito · itálico · sublinhado · cor · link, vindos do editor da dashboard
        # (mesmos segmentos que o e-mail usa → Word e .eml nunca divergem)
        for _segs in parse_intro_lines(_intro):
            p = doc.add_paragraph()
            _intro_para_fmt(p)   # justificado + entrelinha ao menos 16pt (Arial 11)
            for _sg in _segs:
                if not _sg.text:
                    continue
                if _sg.url:
                    _external_hyperlink_run(p, _sg.text, _sg.url, size_pt=11,
                                            color=_sg.color or "0000FF", underline=True,
                                            bold=_sg.bold, italic=_sg.italic)
                else:
                    _run(p, _sg.text, bold=_sg.bold, italic=_sg.italic,
                         underline=_sg.underline, size_pt=11, color_hex=_sg.color or None)
        _blanks(2)   # Mensagem de abertura → Sector Headlines: 2 linhas
    else:
        _blanks(2)   # sem mensagem: logo do Itaú → Sector Headlines: 2 linhas

    # ══════════════════════════════════════════════════════════════════════════
    # SECTOR HEADLINES
    # ══════════════════════════════════════════════════════════════════════════
    _heading_index("Sector Headlines")
    _blanks(1)       # Sector Headlines → notícias: 1 linha

    for sector_key in seen_sectors:
        for item in [it for it in items if it.sector == sector_key]:
            p = doc.add_paragraph()
            _zero_spacing(p)
            _justify(p)
            _add_numPr(p)

            take_color = TAKE_COLOR_HEX.get(item.take)   # "=" → None → preto
            take_sym   = TAKE_SYMBOL.get(item.take, "(=)")

            _run(p, f"{SECTOR_LABEL[sector_key]} -\xa0", bold=True, size_pt=11)   # hífen (headlines), como na referência
            _hyperlink_run(p, item.title, bm_names[item.url])
            # Artigos bilíngues: " \ Título traduzido" vira LINK INTERNO p/ a Free Translation
            if item.translated_title:
                _run(p, " \\ ", size_pt=11)
                _hyperlink_run(p, item.translated_title, bm_tr_names[item.url])
            _run(p, f" [{item.source_name}]", bold=True, size_pt=11)
            _run(p, f" {take_sym}", bold=True, size_pt=11, color_hex=take_color)

    _blanks(2)       # última notícia → Recent Publications: 2 linhas

    # ══════════════════════════════════════════════════════════════════════════
    # RECENT PUBLICATIONS
    # ══════════════════════════════════════════════════════════════════════════
    # helper reusado por Recent Publications e Earnings Review — bullet "SETOR – título (link)"
    def _pub_bullets(pub_list):
        for pub in (pub_list or []):
            name = (pub.get("name") or pub.get("title") or "").strip()
            if not name:
                continue
            sec  = pub.get("sector") if pub.get("sector") in _VALID_SECTORS else "SM"
            link = (pub.get("link") or pub.get("pdf_url") or pub.get("report_url") or "").strip()
            p = doc.add_paragraph(); _zero_spacing(p); _justify(p); _add_numPr(p)
            _run(p, f"{SECTOR_LABEL.get(sec, 'STEEL & MINING')} –\xa0", bold=True, size_pt=11)
            if link:
                _external_hyperlink_run(p, name, link, size_pt=11, color="0000FF")  # azul de link
            else:
                _run(p, name, size_pt=11)

    _heading_index("Recent Publications")
    _blanks(1)       # Recent Publications → publicações: 1 linha
    _recent = config.get("recent_publications") or []
    if _recent:
        _pub_bullets(_recent)
    else:
        p_empty = doc.add_paragraph()
        _zero_spacing(p_empty)
        _run(p_empty, "[Sem publicações — adicione na tela do Clipinator]",
             italic=True, size_pt=9, color_hex="888888")
    # (o gap para a próxima seção é inserido por ela — Preview ou analistas)

    # ══════════════════════════════════════════════════════════════════════════
    # EARNINGS REVIEW / "PREVIEW" (opcional — toggle + nome editável, ex.: "2Q26 Review")
    # ══════════════════════════════════════════════════════════════════════════
    _er = config.get("earnings_review") or {}
    if _er.get("on"):
        _blanks(2)   # última publicação (Recent) → Preview: 2 linhas
        _heading_index((_er.get("label") or "Earnings Review").strip() or "Earnings Review")
        _blanks(1)   # título do Preview → publicações: 1 linha
        _pub_bullets(_er.get("items"))
        _blanks(2)   # última publicação do Preview → analistas: 2 linhas
    else:
        _blanks(2)   # sem Preview: última publicação (Recent) → analistas: 2 linhas

    # ══════════════════════════════════════════════════════════════════════════
    # CONTACTS (analistas configuráveis no admin — fallback = _DEFAULT_ANALYSTS)
    # ══════════════════════════════════════════════════════════════════════════
    _contacts = config.get("analysts") or _DEFAULT_ANALYSTS
    _rendered_any = False
    for _c in _contacts:
        _name  = (_c.get("name")  or "").strip()
        _role  = (_c.get("role")  or "").strip()
        _phone = (_c.get("phone") or "").strip()
        _email = (_c.get("email") or "").strip()
        if not (_name or _email):
            continue
        if _rendered_any:
            _blanks(1)   # entre analistas: 1 linha (conjunto nome→email)
        _rendered_any = True
        if _name:
            p_name = doc.add_paragraph()
            _zero_spacing(p_name)
            _run(p_name, _name, bold=True, size_pt=10, color_hex="FF5000")
        for _line in (_role, _phone):
            if not _line:
                continue
            p_line = doc.add_paragraph()
            _zero_spacing(p_line)
            _run(p_line, _line, size_pt=10)
        if _email:
            p_mail = doc.add_paragraph()
            _zero_spacing(p_mail)
            _mailto_run(p_mail, _email, size_pt=10)   # e-mail = link azul sublinhado (igual à referência)

    _blanks(2)       # último analista → STEEL & MINING: 2 linhas

    # ══════════════════════════════════════════════════════════════════════════
    # CORPOS DOS ARTIGOS POR SETOR
    # ══════════════════════════════════════════════════════════════════════════
    for _si, sector_key in enumerate(seen_sectors):
        sector_items = [it for it in items if it.sector == sector_key]

        if _si > 0:
            _blanks(2)   # entre setores: 2 linhas (a 1ª seção já teve as 2 do último analista → setor)
        _heading_sector(SECTOR_LABEL[sector_key])
        _blanks(1)       # cabeçalho do setor → 1º artigo: 1 linha

        for _ii, item in enumerate(sector_items):
            if _ii > 0:
                _blanks(2)   # entre notícias: 2 linhas em branco (do último parágrafo → título da próxima)
            bm_name = bm_names[item.url]

            # ── Título: bold, highlight amarelo, bookmarked ───────────────────
            # Título/Source/Corpo = conjunto tight (sem espaçamento interno). Parágrafos do
            # corpo são separados por 1 linha em branco (em _render_blocks); entre notícias, 2.
            # Para artigos bilíngues: sufixo "(Original)" no título.
            p_title = doc.add_paragraph()
            _zero_spacing(p_title)
            p_title.paragraph_format.space_before = Pt(0)
            p_title.paragraph_format.space_after  = Pt(2)
            _justify(p_title)
            title_display = f"{item.title} (Original)" if item.translated_title else item.title
            _run(p_title, title_display, bold=True, size_pt=12, hl="yellow")
            _add_bookmark(p_title, bm_name)

            # ── Source ────────────────────────────────────────────────────────
            p_src = doc.add_paragraph()
            _zero_spacing(p_src)
            _justify(p_src)
            _run(p_src, f"Source: {item.source_name}",
                 italic=True, size_pt=12)   # 12pt itálico PRETO, sem espaçamento (igual à referência)

            # ── Corpo do artigo ───────────────────────────────────────────────
            blocks = _html_to_blocks(item.body) if item.body else []

            def _body_para(size_pt=9, bold=False, italic=False,
                           color_hex=None, indent_dxa=0,
                           space_before_pt=0, space_after_pt=0):
                """Cria parágrafo de corpo TIGHT (space_after=0). A separação de 1 linha entre os
                parágrafos do scraping é feita por LINHAS EM BRANCO reais em _render_blocks
                (o usuário quer 'pular uma linha' de verdade, não só espaçamento)."""
                p = doc.add_paragraph()
                pPr = p._p.get_or_add_pPr()
                sp = OxmlElement("w:spacing")
                sp.set(qn("w:before"), str(int(space_before_pt * 20)))
                sp.set(qn("w:beforeAutospacing"), "0")
                sp.set(qn("w:after"),  str(int(space_after_pt * 20)))
                sp.set(qn("w:afterAutospacing"),  "0")
                sp.set(qn("w:line"),      "240")
                sp.set(qn("w:lineRule"), "auto")
                pPr.append(sp)
                _justify(p)
                if indent_dxa:
                    ind = OxmlElement("w:ind")
                    ind.set(qn("w:left"), str(indent_dxa))
                    pPr.append(ind)
                return p

            def _render_blocks(blocks_list: list[dict], item_domain: str) -> None:
                """Renderiza os blocos como parágrafos Word, com 1 LINHA EM BRANCO real entre
                cada bloco (parágrafos do scraping não viram texto corrido)."""
                # Platts: os "bullets" do corpo eram os highlights (resumo em negrito), que o
                # usuário pediu p/ REMOVER. O corpo do Platts não tem lista real → descarta os
                # list_item ANTES de renderizar (conserta também corpos já cacheados com highlights).
                if item_domain == "core.spglobal.com":
                    blocks_list = [b for b in blocks_list if b.get("type") != "list_item"]
                if not blocks_list:
                    p = doc.add_paragraph()
                    _zero_spacing(p)
                    _run(p, "[Corpo do artigo não disponível]",
                         italic=True, size_pt=9, color_hex="888888")
                    return
                _prev = None
                for block in blocks_list:
                    btype = block["type"]
                    # 1 linha em branco entre blocos — exceto antes do 1º e entre bullets consecutivos
                    if _prev is not None and not (_prev == "list_item" and btype == "list_item"):
                        _blanks(1)
                    _prev = btype
                    if btype == "text":
                        p = _body_para()
                        _run(p, block["text"], size_pt=9, bold=block.get("bold", False))
                    elif btype == "h3":
                        p = _body_para(space_before_pt=6, space_after_pt=2)
                        _run(p, block["text"], bold=True, size_pt=10.5)
                    elif btype == "h4":
                        p = _body_para(space_before_pt=4, space_after_pt=2)
                        _run(p, block["text"], bold=True, size_pt=10)
                    elif btype == "blockquote":
                        p = _body_para(indent_dxa=200, space_before_pt=4, space_after_pt=4)
                        _run(p, block["text"], italic=True, size_pt=9, color_hex="595959")
                    elif btype == "list_item":
                        if item_domain == "core.spglobal.com":
                            p = _body_para()
                            _run(p, block["text"], size_pt=9)
                        else:
                            p = _body_para(space_after_pt=2)
                            pPr = p._p.get_or_add_pPr()
                            ind = OxmlElement("w:ind")
                            ind.set(qn("w:left"),    "280")
                            ind.set(qn("w:hanging"), "180")
                            pPr.append(ind)
                            _run(p, "•\xa0", bold=True, size_pt=9)
                            _run(p, block["text"], size_pt=9)
                    elif btype == "image":
                        _add_image_para(block["src"], item.url)

            # Renderiza corpo original
            _render_blocks(blocks, item.domain)

            # ── Free Translation (artigos bilíngues) ──────────────────────────
            if item.translated_title:
                _blanks(2)   # original → tradução: 2 linhas (mesmo espaçamento de entre notícias)
                # Título da tradução
                p_tr_title = doc.add_paragraph()
                _zero_spacing(p_tr_title)
                p_tr_title.paragraph_format.space_before = Pt(6)
                p_tr_title.paragraph_format.space_after  = Pt(2)
                _justify(p_tr_title)
                _run(p_tr_title,
                     f"{item.translated_title} (Free Translation)",
                     bold=True, size_pt=12, hl="yellow")
                _add_bookmark(p_tr_title, bm_tr_names[item.url])   # alvo do link interno da headline

                # Source da tradução
                p_tr_src = doc.add_paragraph()
                _zero_spacing(p_tr_src)
                _justify(p_tr_src)
                _run(p_tr_src, f"Source: {item.source_name}",
                     italic=True, size_pt=12)

                # Corpo traduzido
                trans_blocks = _html_to_blocks(item.translated_body) if item.translated_body else []
                _render_blocks(trans_blocks, item.domain)

            # Sem separador extra entre artigos: o space_after (~1 linha) do último parágrafo
            # do corpo já dá a separação de 1 linha para o próximo artigo (ou p/ o próximo setor).

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Public API ────────────────────────────────────────────────────────────────

def build_docx(items: "list[ClippingItem]", d: date | None = None, config: dict | None = None) -> bytes:
    """Gera o .docx (bytes) a partir de ClippingItems JÁ com corpo preenchido.

    config (opcional): {intro:{on,text}, recent_publications:[{name,sector,link}],
    earnings_review:{on,label,items:[{name,sector,link}]}} — vindo da tela do Clipinator.
    """
    return _build_word(list(items), d or date.today(), config)


def domain_supported(url: str) -> bool:
    """Sempre True — todos os domínios suportados para clipping."""
    return True


def generate_clipping(
    urls:    list[str],
    takes:   list[str],
    sectors: list[str] | None = None,
    d:       date | None = None,
) -> tuple[Path | None, list[tuple[str, str]]]:
    """Gera Word .docx no formato Itaú BBA. Retorna (path, erros).

    urls, takes e sectors devem ter o mesmo comprimento e mesma ordem.
    sectors: valor explícito ('SM','PP','NR','CEMENT') ou '' para auto-detect.
    """
    from .store import get_article

    d       = d or date.today()
    sectors = sectors or []
    items:  list[ClippingItem] = []
    errors: list[tuple[str, str]] = []

    for i, url in enumerate(urls):
        take        = takes[i]   if i < len(takes)   else "="
        user_sector = sectors[i] if i < len(sectors) else ""
        article     = get_article(url)
        if article is None:
            errors.append((url, "artigo não encontrado no banco de dados"))
            continue

        sector = (
            user_sector
            if user_sector in _VALID_SECTORS
            else detect_sector(article.domain, article.matched_keywords, article.title)
        )

        items.append(ClippingItem(
            url=url,
            title=article.title,
            source_name=article.source_name,
            body=article.body or "",
            matched_keywords=article.matched_keywords,
            domain=article.domain,
            take=take,
            sector=sector,
        ))

    if not items:
        return None, errors

    # ── Auto-fetch de corpo para artigos sem body ─────────────────────────────
    # Artigos de fontes regulares (mining.com, el financiero) são indexados
    # sem corpo — só snippet. Se o usuário não os abriu no leitor antes, body
    # fica vazio e o Word mostraria "[Corpo do artigo não disponível]".
    # Aqui buscamos o corpo em paralelo para todos esses artigos, salvamos no
    # DB e preenchemos o ClippingItem antes de gerar o Word.
    from concurrent.futures import ThreadPoolExecutor, as_completed
    from .store import save_article_body

    needs_body = [it for it in items if not it.body and it.domain not in _SKIP_AUTO_FETCH]
    if needs_body:
        log.info(
            "clipping: auto-fetch de corpo para %d artigo(s) sem body: %s",
            len(needs_body),
            [it.source_name for it in needs_body],
        )

        def _fetch_for_item(it: ClippingItem) -> tuple[ClippingItem, str]:
            return it, _fetch_body_regular(it.url)

        with ThreadPoolExecutor(max_workers=4, thread_name_prefix="clip_fetch") as _ex:
            for fut in as_completed({_ex.submit(_fetch_for_item, it): it for it in needs_body}):
                try:
                    it, body = fut.result()
                    if body:
                        it.body = body
                        try:
                            save_article_body(it.url, body)
                        except Exception:
                            pass
                        log.info(
                            "clipping: body auto-fetchado para '%s' (%d chars)",
                            it.title[:60], len(body),
                        )
                    else:
                        log.debug("clipping: auto-fetch sem resultado para '%s'", it.title[:60])
                except Exception as e:
                    log.debug("clipping: erro no auto-fetch: %s", e)

    # ── Tradução para artigos bilíngues ──────────────────────────────────────────
    # Valor, Estadão e El Financiero saem em dois blocos: original + Free Translation.
    # A tradução é feita em paralelo (máx 3 threads) para não bloquear muito.
    bilingual_items = [it for it in items if it.domain in _BILINGUAL_DOMAINS]
    if bilingual_items:
        log.info("clipping: traduzindo %d artigo(s) bilíngue(s) para inglês...", len(bilingual_items))

        def _translate_item(it: ClippingItem) -> tuple[ClippingItem, str, str]:
            lang = _DOMAIN_LANG.get(it.domain, "Portuguese")
            t_title, t_body = _translate_to_english(it.title, it.body, lang)
            return it, t_title, t_body

        with ThreadPoolExecutor(max_workers=3, thread_name_prefix="clip_transl") as _tex:
            futures = {_tex.submit(_translate_item, it): it for it in bilingual_items}
            for fut in as_completed(futures):
                try:
                    it, t_title, t_body = fut.result()
                    if t_title:
                        it.translated_title = t_title
                        it.translated_body  = t_body
                        log.info("clipping: traduzido '%s' → '%s'", it.title[:55], t_title[:55])
                    else:
                        log.debug("clipping: tradução sem resultado para '%s'", it.title[:55])
                except Exception as e:
                    log.debug("clipping: erro na tradução: %s", e)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    filename = f"{clipping_basename(d)}.docx"
    out_path = OUT_DIR / filename

    try:
        docx_bytes = _build_word(items, d)
        out_path.write_bytes(docx_bytes)
    except Exception as e:
        log.exception("Erro ao gerar Word: %s", e)
        return None, errors + [(str(out_path), str(e))]

    return out_path, errors


# Legacy alias
def generate_eml(urls: list[str], d: date | None = None) -> tuple[Path | None, list[tuple[str, str]]]:
    takes = ["="] * len(urls)
    return generate_clipping(urls, takes, d)
